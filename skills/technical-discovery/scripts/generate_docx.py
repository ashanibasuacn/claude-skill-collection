#!/usr/bin/env python3
"""
generate_docx.py — Convert one or more discovery markdown files into a formal
Current Landscape / Technical Discovery DOCX (neutral, client-agnostic styling).

Usage:
  python generate_docx.py --source <md1> [<md2> ...] --output <docx> --plan <json>

The plan JSON carries metadata, a section_mapping (discovery-section key -> exact
"## Heading" taken from the source), a tbd_sections list, and the source_files
list used to build the Source Inventory appendix.
"""

import argparse
import json
import re
import sys
import subprocess
from datetime import date
from pathlib import Path

# Auto-install python-docx if not present
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...", flush=True)
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Markdown parsing
# ─────────────────────────────────────────────

def parse_md_sections(text: str) -> dict:
    """Return {heading_title: body_text} for every heading in the markdown."""
    sections = {}
    current_title = None
    current_body = []
    for line in text.splitlines():
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            if current_title is not None:
                sections[current_title] = '\n'.join(current_body).strip()
            current_title = m.group(2).strip()
            current_body = []
        else:
            if current_title is not None:
                current_body.append(line)
    if current_title is not None:
        sections[current_title] = '\n'.join(current_body).strip()
    return sections


def find_content(sections: dict, keys) -> str | None:
    if isinstance(keys, str):
        keys = [keys]
    for title, body in sections.items():
        for key in keys:
            if key.lower() in title.lower():
                return body if body else None
    return None


def content_from_plan(sections: dict, plan: dict, mapping_key: str) -> str | None:
    heading = plan.get('section_mapping', {}).get(mapping_key)
    if not heading:
        return None
    title = re.sub(r'^#{1,4}\s+', '', heading).strip()
    return sections.get(title) or None


def get_section(sections: dict, plan: dict, mapping_key: str, fallback_keys=None) -> str | None:
    content = content_from_plan(sections, plan, mapping_key)
    if content:
        return content
    if fallback_keys:
        return find_content(sections, fallback_keys)
    return None


# ─────────────────────────────────────────────
# DOCX helpers (neutral styling)
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def bold_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(name):
        run = p.add_run()
        for tag in ('begin', 'instrText', 'end'):
            el = OxmlElement('w:fldChar' if tag != 'instrText' else 'w:instrText')
            if tag != 'instrText':
                el.set(qn('w:fldCharType'), tag)
            else:
                el.text = name
            run._r.append(el)

    p.add_run('Page ')
    _field('PAGE')
    p.add_run(' of ')
    _field('NUMPAGES')


def add_doc_header(doc, title: str, version: str, status: str):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{title}   v{version} — {status}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def tbd(doc, label: str):
    """Greyed-out placeholder for a section with no source content."""
    p = doc.add_paragraph()
    run = p.add_run(f"[TBD — not covered in the source material: {label}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def strip_inline_md(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def render_md_body(doc, content: str):
    if not content:
        return
    lines = content.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        if line.strip().startswith('|') and '|' in line:
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            _render_md_table(doc, tbl_lines)
            continue
        h4 = re.match(r'^####\s+(.+)$', line)
        h3 = re.match(r'^###\s+(.+)$', line)
        h2 = re.match(r'^##\s+(.+)$', line)
        if h4:
            doc.add_heading(strip_inline_md(h4.group(1)), level=4)
        elif h3:
            doc.add_heading(strip_inline_md(h3.group(1)), level=4)
        elif h2:
            doc.add_heading(strip_inline_md(h2.group(1)), level=4)
        elif re.match(r'^[-*]\s+', line):
            doc.add_paragraph(strip_inline_md(re.sub(r'^[-*]\s+', '', line)), style='List Bullet')
        elif re.match(r'^\d+\.\s+', line):
            doc.add_paragraph(strip_inline_md(re.sub(r'^\d+\.\s+', '', line)), style='List Number')
        elif line.startswith('> '):
            doc.add_paragraph(strip_inline_md(line[2:]), style='Quote')
        elif line.strip():
            doc.add_paragraph(strip_inline_md(line))
        i += 1


def _render_md_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r'^\s*\|?[-: |]+\|?\s*$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'
    for r_i, row_data in enumerate(rows):
        for c_i, text in enumerate(row_data[:cols]):
            cell = tbl.rows[r_i].cells[c_i]
            cell.text = strip_inline_md(text)
            if r_i == 0:
                set_cell_bg(cell, 'E8E8E8')
                bold_cell(cell)


def parse_glossary_entries(content: str) -> list[tuple[str, str]]:
    entries = []
    for line in content.splitlines():
        m = re.match(r'^\*\*(.+?)\*\*[:\s]+(.+)$', line)
        if m:
            entries.append((m.group(1).strip(), m.group(2).strip()))
            continue
        m = re.match(r'^[-*]\s+(.+?):\s+(.+)$', line)
        if m:
            entries.append((m.group(1).strip(), m.group(2).strip()))
            continue
        if '|' in line and not re.match(r'^\s*\|?[-: |]+\|?\s*$', line):
            parts = [p.strip() for p in line.strip('|').split('|')]
            if len(parts) >= 2 and parts[0] and parts[1]:
                entries.append((parts[0], parts[1]))
    return [(t, d) for t, d in entries
            if t.lower() not in ('term', 'acronym', 'name', 'word', 'definition')]


def add_kv_table(doc, entries, col_header='Term', empty_note='No entries found in source'):
    tbl = doc.add_table(rows=1 + max(len(entries), 0), cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0]
    hdr.cells[0].text = col_header
    hdr.cells[1].text = 'Definition'
    for cell in hdr.cells:
        set_cell_bg(cell, 'D4D4D4')
        bold_cell(cell)
    for i, (term, definition) in enumerate(entries):
        row = tbl.rows[i + 1]
        row.cells[0].text = term
        row.cells[1].text = definition
        bold_cell(row.cells[0])
    if not entries:
        row = tbl.add_row()
        row.cells[0].text = '—'
        row.cells[1].text = empty_note


# ─────────────────────────────────────────────
# Document generator — Current Landscape / Technical Discovery
# ─────────────────────────────────────────────

def section(doc, sections, plan, number, title, mapping_key, fallbacks):
    """Render a numbered section: mapped content, else a TBD placeholder."""
    doc.add_heading(f'{number} {title}', level=1 if '.' not in number else 2)
    c = get_section(sections, plan, mapping_key, fallbacks)
    render_md_body(doc, c) if c else tbd(doc, f'{number} {title}')


def generate(plan: dict, source_text: str, source_files: list[str], output_path: Path):
    doc = Document()

    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
            setattr(sec, attr, Cm(2.54))

    title = plan.get('title', 'Current Landscape — Technical Discovery')
    version = plan.get('version', '1.0')
    status = plan.get('status', 'Draft')
    project = plan.get('project', '')
    client = plan.get('client', '')
    author = plan.get('author', '')
    doc_date = plan.get('date', str(date.today()))
    confidentiality = plan.get('confidentiality', 'Internal')

    sections = parse_md_sections(source_text)

    add_doc_header(doc, title, version, status)
    add_page_number_footer(doc)

    # ── COVER PAGE ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if project:
        p2 = doc.add_paragraph(project)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.size = Pt(14)
    doc.add_paragraph()
    meta = doc.add_table(rows=6, cols=2)
    meta.style = 'Table Grid'
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate([
        ('Document Title', title),
        ('Project', project or '—'),
        ('Client / Organisation', client or '—'),
        ('Author', author or '—'),
        ('Version', f'v{version} — {status}'),
        ('Date', doc_date),
    ]):
        row = meta.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], 'F0F0F0')
        bold_cell(row.cells[0])
    doc.add_paragraph()
    cp = doc.add_paragraph(f'Classification: {confidentiality}')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.add_page_break()

    # ── DOCUMENT CONTROL ──
    doc.add_heading('Document Control', level=1)
    doc.add_heading('Version History', level=2)
    vh = doc.add_table(rows=2, cols=5)
    vh.style = 'Table Grid'
    for i, h_label in enumerate(['Version', 'Date', 'Author', 'Status', 'Description of Changes']):
        cell = vh.rows[0].cells[i]
        cell.text = h_label
        set_cell_bg(cell, 'D4D4D4')
        bold_cell(cell)
    for i, val in enumerate([f'v{version}', doc_date, author or '—', status, 'Initial version']):
        vh.rows[1].cells[i].text = val
    doc.add_page_break()

    # ── TABLE OF CONTENTS ──
    doc.add_heading('Table of Contents', level=1)
    add_toc(doc)
    p = doc.add_paragraph()
    run = p.add_run('Open this document in Microsoft Word and press Ctrl+A → F9 to update the TOC.')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ──
    doc.add_heading('1. Executive Summary', level=1)
    c = get_section(sections, plan, 'executive_summary',
                    ['executive summary', 'summary', 'overview', 'introduction', 'background'])
    render_md_body(doc, c) if c else tbd(doc, '1. Executive Summary')

    # ── 2. DISCOVERY SCOPE & APPROACH ──
    doc.add_heading('2. Discovery Scope & Approach', level=1)
    section(doc, sections, plan, '2.1', 'Scope', 'scope',
            ['scope', 'in scope', 'objectives', 'purpose'])
    section(doc, sections, plan, '2.2', 'Approach & Sources', 'approach',
            ['approach', 'method', 'methodology', 'sources', 'discovery'])
    section(doc, sections, plan, '2.3', 'Assumptions', 'assumptions',
            ['assumptions', 'dependencies'])

    # ── 3. BUSINESS CONTEXT ──
    doc.add_heading('3. Business Context', level=1)
    section(doc, sections, plan, '3.1', 'Business Capabilities & Domain Overview', 'business_capabilities',
            ['business capabilities', 'business context', 'domain', 'business overview', 'capabilities'])
    section(doc, sections, plan, '3.2', 'Business & Process Flows', 'business_flows',
            ['business flow', 'process flow', 'workflows', 'user journey', 'processes', 'flows'])
    section(doc, sections, plan, '3.3', 'Users & Stakeholders', 'users_stakeholders',
            ['users', 'user classes', 'stakeholders', 'actors', 'personas', 'roles'])

    # ── 4. APPLICATION & ARCHITECTURE LANDSCAPE ──
    doc.add_heading('4. Application & Architecture Landscape', level=1)
    section(doc, sections, plan, '4.1', 'System & Application Inventory', 'system_inventory',
            ['application inventory', 'system inventory', 'applications', 'systems', 'components inventory'])
    section(doc, sections, plan, '4.2', 'Architecture Overview', 'architecture_overview',
            ['architecture', 'solution architecture', 'system architecture', 'component', 'logical view'])
    section(doc, sections, plan, '4.3', 'Integration Landscape', 'integration_landscape',
            ['integration', 'interfaces', 'apis', 'api', 'messaging', 'external systems'])

    # ── 5. TECHNOLOGY STACK ──
    doc.add_heading('5. Technology Stack', level=1)
    c = get_section(sections, plan, 'technology_stack',
                    ['technology stack', 'tech stack', 'technologies', 'languages', 'frameworks', 'runtime'])
    render_md_body(doc, c) if c else tbd(doc, '5. Technology Stack')

    # ── 6. TOOLS & PLATFORMS ──
    doc.add_heading('6. Tools & Platforms', level=1)
    section(doc, sections, plan, '6.1', 'Development & Delivery (Build / CI-CD)', 'dev_delivery',
            ['ci/cd', 'cicd', 'build', 'pipeline', 'devops', 'development tools', 'tooling'])
    section(doc, sections, plan, '6.2', 'Hosting & Infrastructure', 'hosting_infra',
            ['hosting', 'infrastructure', 'platform', 'cloud', 'deployment', 'environment'])
    section(doc, sections, plan, '6.3', 'Monitoring & Observability', 'monitoring',
            ['monitoring', 'observability', 'logging', 'alerting', 'telemetry'])

    # ── 7. DATA LANDSCAPE ──
    doc.add_heading('7. Data Landscape', level=1)
    section(doc, sections, plan, '7.1', 'Data Model', 'data_model',
            ['data model', 'entities', 'schema', 'erd', 'domain model'])
    section(doc, sections, plan, '7.2', 'Data Stores & Flows', 'data_stores',
            ['data store', 'databases', 'data flow', 'data sources', 'storage', 'data'])

    # ── 8. OPERATIONAL & NON-FUNCTIONAL PROFILE ──
    doc.add_heading('8. Operational & Non-Functional Profile', level=1)
    section(doc, sections, plan, '8.1', 'Security Posture', 'security',
            ['security', 'authentication', 'authorization', 'compliance', 'iam'])
    section(doc, sections, plan, '8.2', 'Availability & Resilience', 'availability',
            ['availability', 'resilience', 'reliability', 'disaster recovery', 'sla', 'uptime'])
    section(doc, sections, plan, '8.3', 'Performance & Scalability', 'performance',
            ['performance', 'scalability', 'capacity', 'throughput', 'latency'])

    # ── 9. KNOWN ISSUES, PAIN POINTS & TECHNICAL DEBT ──
    doc.add_heading('9. Known Issues, Pain Points & Technical Debt', level=1)
    c = get_section(sections, plan, 'known_issues',
                    ['known issues', 'pain points', 'technical debt', 'issues', 'problems', 'gaps', 'limitations'])
    render_md_body(doc, c) if c else tbd(doc, '9. Known Issues, Pain Points & Technical Debt')

    # ── 10. CONSTRAINTS, RISKS & DEPENDENCIES ──
    doc.add_heading('10. Constraints, Risks & Dependencies', level=1)
    c = get_section(sections, plan, 'constraints_risks',
                    ['constraints', 'risks', 'dependencies', 'risk'])
    render_md_body(doc, c) if c else tbd(doc, '10. Constraints, Risks & Dependencies')

    # ── 11. OBSERVATIONS & RECOMMENDATIONS ──
    doc.add_heading('11. Observations & Recommendations', level=1)
    c = get_section(sections, plan, 'recommendations',
                    ['recommendations', 'observations', 'next steps', 'opportunities', 'findings'])
    render_md_body(doc, c) if c else tbd(doc, '11. Observations & Recommendations')

    # ── APPENDIX A: GLOSSARY ──
    doc.add_page_break()
    doc.add_heading('Appendix A: Glossary', level=1)
    c = get_section(sections, plan, 'glossary', ['glossary', 'definitions', 'terms and definitions', 'terms'])
    if c:
        entries = parse_glossary_entries(c)
        add_kv_table(doc, entries, 'Term') if entries else render_md_body(doc, c)
    else:
        add_kv_table(doc, [], 'Term')

    # ── APPENDIX B: ACRONYMS ──
    doc.add_heading('Appendix B: Acronyms and Abbreviations', level=1)
    c = get_section(sections, plan, 'acronyms', ['acronyms', 'abbreviations'])
    entries = parse_glossary_entries(c) if c else []
    add_kv_table(doc, entries, 'Acronym')

    # ── APPENDIX C: SOURCE INVENTORY ──
    doc.add_heading('Appendix C: Source Inventory', level=1)
    doc.add_paragraph('This discovery document was assembled from the following source materials:')
    si = doc.add_table(rows=1 + max(len(source_files), 0), cols=2)
    si.style = 'Table Grid'
    si.rows[0].cells[0].text = '#'
    si.rows[0].cells[1].text = 'Source file'
    for cell in si.rows[0].cells:
        set_cell_bg(cell, 'D4D4D4')
        bold_cell(cell)
    if source_files:
        for i, fname in enumerate(source_files):
            si.rows[i + 1].cells[0].text = str(i + 1)
            si.rows[i + 1].cells[1].text = fname
    else:
        r = si.add_row()
        r.cells[0].text = '—'
        r.cells[1].text = 'No source files recorded'

    # ── APPENDIX D: OPEN QUESTIONS ──
    c = get_section(sections, plan, 'open_questions',
                    ['open questions', 'open issues', 'tbd', 'to be confirmed', 'questions'])
    doc.add_heading('Appendix D: Open Questions', level=1)
    if c:
        render_md_body(doc, c)
    else:
        doc.add_paragraph('Open questions surfaced during discovery should be captured here as the '
                          'assessment progresses. None were recorded in the source material.')

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Document saved: {output_path}")


# ─────────────────────────────────────────────
# CLI entry point
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Generate a Technical Discovery DOCX from markdown sources')
    parser.add_argument('--source', required=True, nargs='+', help='One or more source markdown files')
    parser.add_argument('--output', required=True, help='Path for output DOCX file')
    parser.add_argument('--plan',   required=True, help='Path to document plan JSON')
    args = parser.parse_args()

    output = Path(args.output)
    plan_file = Path(args.plan)
    if not plan_file.exists():
        sys.exit(f'Plan file not found: {plan_file}')
    plan = json.loads(plan_file.read_text(encoding='utf-8'))

    combined = []
    source_files = []
    for sp in args.source:
        p = Path(sp)
        if not p.exists():
            sys.exit(f'Source file not found: {p}')
        source_files.append(p.name)
        combined.append(f"\n\n# [SOURCE: {p.name}]\n\n" + p.read_text(encoding='utf-8', errors='replace'))

    # plan may override the recorded source file names (e.g. friendly labels)
    recorded = plan.get('source_files') or source_files
    generate(plan, "\n".join(combined), recorded, output)


if __name__ == '__main__':
    main()
