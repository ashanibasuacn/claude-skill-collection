#!/usr/bin/env python3
"""
generate_docx.py — Convert requirements MD to IEEE 830 SRS DOCX.

Usage:
  python generate_docx.py --source <md_file> --output <docx_file> --plan <json_file>
"""

import argparse
import json
import re
import sys
import subprocess
from datetime import date
from pathlib import Path

# Auto-install python-docx if not present
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...", flush=True)
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Markdown parsing
# ─────────────────────────────────────────────

def parse_md_sections(text: str) -> dict:
    """Return {heading_title: immediate_body_text} for every heading (flat, no nesting)."""
    sections = {}
    current_title = None
    current_body = []

    for line in text.splitlines():
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            if current_title is not None:
                sections[current_title] = '\n'.join(current_body).strip()
            current_title = m.group(2).strip()
            current_body = []
        else:
            if current_title is not None:
                current_body.append(line)

    if current_title is not None:
        sections[current_title] = '\n'.join(current_body).strip()

    return sections


def get_section_full(source_text: str, heading_ref: str) -> str | None:
    """
    Return ALL content under a heading — including every nested sub-heading and its
    body — until the next heading at the SAME OR HIGHER level.

    heading_ref may include leading # markers (e.g. '## Part 1: Functional Requirements')
    or just the bare title text.
    """
    clean = re.sub(r'^#{1,4}\s+', '', heading_ref).strip()
    lines = source_text.splitlines()

    target_level: int | None = None
    start_idx: int | None = None

    for i, line in enumerate(lines):
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m and m.group(2).strip() == clean:
            target_level = len(m.group(1))
            start_idx = i + 1
            break

    if start_idx is None:
        return None

    body: list[str] = []
    for line in lines[start_idx:]:
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m and len(m.group(1)) <= target_level:
            break
        body.append(line)

    result = '\n'.join(body).strip()
    return result if result else None


def find_content(sections: dict, keys) -> str | None:
    """Return body text for the first heading whose title contains any key (case-insensitive)."""
    if isinstance(keys, str):
        keys = [keys]
    for title, body in sections.items():
        for key in keys:
            if key.lower() in title.lower():
                return body if body else None
    return None


def find_content_full(source_text: str, keys) -> str | None:
    """
    Like find_content but returns the FULL hierarchical content (via get_section_full)
    for the first matching heading title.
    """
    if isinstance(keys, str):
        keys = [keys]
    for line in source_text.splitlines():
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            title = m.group(2).strip()
            for key in keys:
                if key.lower() in title.lower():
                    content = get_section_full(source_text, title)
                    if content:
                        return content
    return None


# ─────────────────────────────────────────────
# DOCX helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def bold_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def add_toc(doc):
    """Insert a Word TOC field (auto-updates when opened in Word)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)


def add_page_number_footer(doc):
    """Add centred 'Page X of Y' footer."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _fld(name):
        run = p.add_run()
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(begin)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {name} '
        run._r.append(instr)
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        run._r.append(end)

    p.add_run('Page ')
    _fld('PAGE')
    p.add_run(' of ')
    _fld('NUMPAGES')


def add_doc_header(doc, title: str, version: str, status: str):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{title}   v{version} — {status}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def tbd(doc, label: str):
    """Insert a greyed-out TBD placeholder."""
    p = doc.add_paragraph()
    run = p.add_run(f"[TBD — {label}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def strip_inline_md(text: str) -> str:
    """Strip common inline markdown markers from a string."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def render_md_body(doc, content: str):
    """Render a block of markdown text into Word paragraphs."""
    if not content:
        return

    lines = content.splitlines()
    i = 0
    in_code = False

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules — they are MD structural separators, not prose
        if re.match(r'^-{3,}\s*$', line) or re.match(r'^\*{3,}\s*$', line):
            i += 1
            continue

        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue

        # Markdown table
        if line.strip().startswith('|') and '|' in line:
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            _render_md_table(doc, tbl_lines)
            continue

        # Blockquote — render as indented italic paragraph
        if line.startswith('> '):
            inner = line[2:].strip()
            # Strip bold markers from blockquote lead text
            inner = re.sub(r'\*\*(.+?)\*\*', r'\1', inner)
            inner = re.sub(r'\*(.+?)\*', r'\1', inner)
            p = doc.add_paragraph(style='Quote')
            p.add_run(inner)
            i += 1
            continue

        # Subheadings (inside a section body)
        h4 = re.match(r'^####\s+(.+)$', line)
        h3 = re.match(r'^###\s+(.+)$', line)
        h2 = re.match(r'^##\s+(.+)$', line)
        if h4:
            doc.add_heading(strip_inline_md(h4.group(1)), level=4)
        elif h3:
            doc.add_heading(strip_inline_md(h3.group(1)), level=3)
        elif h2:
            doc.add_heading(strip_inline_md(h2.group(1)), level=2)
        elif re.match(r'^[-*]\s+', line):
            doc.add_paragraph(strip_inline_md(re.sub(r'^[-*]\s+', '', line)), style='List Bullet')
        elif re.match(r'^\d+\.\s+', line):
            doc.add_paragraph(strip_inline_md(re.sub(r'^\d+\.\s+', '', line)), style='List Number')
        elif line.strip():
            doc.add_paragraph(strip_inline_md(line))

        i += 1


def _render_md_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r'^\s*\|?[-: |]+\|?\s*$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'

    for r_i, row_data in enumerate(rows):
        for c_i, text in enumerate(row_data[:cols]):
            cell = tbl.rows[r_i].cells[c_i]
            cell.text = strip_inline_md(text)
            if r_i == 0:
                set_cell_bg(cell, 'E8E8E8')
                bold_cell(cell)


def parse_glossary_entries(content: str) -> list[tuple[str, str]]:
    """Extract (term, definition) pairs from glossary markdown."""
    entries = []
    for line in content.splitlines():
        m = re.match(r'^\*\*(.+?)\*\*[:\s]+(.+)$', line)
        if m:
            entries.append((m.group(1).strip(), m.group(2).strip()))
            continue
        m = re.match(r'^[-*]\s+(.+?):\s+(.+)$', line)
        if m:
            entries.append((m.group(1).strip(), m.group(2).strip()))
            continue
        if '|' in line and not re.match(r'^\s*\|?[-: |]+\|?\s*$', line):
            parts = [p.strip() for p in line.strip('|').split('|')]
            if len(parts) >= 2 and parts[0] and parts[1]:
                entries.append((parts[0], parts[1]))

    skip = {'term', 'acronym', 'name', 'word', 'definition', 'actor', 'description',
            'gap id', 'area', 'open question', 'id', 'category', 'requirement',
            'workflow', 'web portal', 'mobile app', 'notes'}
    return [(strip_inline_md(t), strip_inline_md(d))
            for t, d in entries
            if t.lower() not in skip and len(t) < 80]


def add_glossary_table(doc, entries: list[tuple[str, str]], col_header='Term'):
    tbl = doc.add_table(rows=1 + len(entries), cols=2)
    tbl.style = 'Table Grid'

    hdr = tbl.rows[0]
    hdr.cells[0].text = col_header
    hdr.cells[1].text = 'Definition'
    for cell in hdr.cells:
        set_cell_bg(cell, 'D4D4D4')
        bold_cell(cell)

    for i, (term, definition) in enumerate(entries):
        row = tbl.rows[i + 1]
        row.cells[0].text = term
        row.cells[1].text = definition
        bold_cell(row.cells[0])

    if not entries:
        row = tbl.add_row()
        row.cells[0].text = '—'
        row.cells[1].text = 'No entries defined in source'


# Common acronyms auto-populated when the source MD has no dedicated acronym section.
# Add project-specific entries via section_mapping["acronyms"] in the plan JSON.
_COMMON_ACRONYMS = [
    ('API',    'Application Programming Interface'),
    ('AWS',    'Amazon Web Services'),
    ('BFF',    'Backend for Frontend'),
    ('CI/CD',  'Continuous Integration / Continuous Delivery'),
    ('CIP',    'Common Identity Provider (Manheim Auth Token)'),
    ('CRUD',   'Create, Read, Update, Delete'),
    ('CSV',    'Comma-Separated Values'),
    ('DL',     "Driver's Licence"),
    ('DR',     'Disaster Recovery'),
    ('FR',     'Functional Requirement'),
    ('FG',     'Functional Gap'),
    ('IAM',    'Identity and Access Management'),
    ('JWT',    'JSON Web Token'),
    ('KMS',    'Key Management Service (AWS)'),
    ('MFA',    'Multi-Factor Authentication'),
    ('NFR',    'Non-Functional Requirement'),
    ('NFG',    'Non-Functional Gap'),
    ('OIDC',   'OpenID Connect'),
    ('PII',    'Personally Identifiable Information'),
    ('RBAC',   'Role-Based Access Control'),
    ('RPO',    'Recovery Point Objective'),
    ('RTO',    'Recovery Time Objective'),
    ('S3',     'Amazon Simple Storage Service'),
    ('SAST',   'Static Application Security Testing'),
    ('SCA',    'Software Composition Analysis'),
    ('SLA',    'Service Level Agreement'),
    ('SLO',    'Service Level Objective'),
    ('SRS',    'Software Requirements Specification'),
    ('SSO',    'Single Sign-On'),
    ('UKG',    'UKG — workforce management platform (formerly Kronos)'),
    ('UTC',    'Coordinated Universal Time'),
    ('WCAG',   'Web Content Accessibility Guidelines'),
]


# ─────────────────────────────────────────────
# Main document generator
# ─────────────────────────────────────────────

def generate(plan: dict, source_text: str, output_path: Path):
    doc = Document()

    # Page layout: A4, 2.54 cm margins
    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
            setattr(sec, attr, Cm(2.54))

    title          = plan.get('title', 'Software Requirements Specification')
    version        = plan.get('version', '1.0')
    status         = plan.get('status', 'Draft')
    project        = plan.get('project', '')
    client         = plan.get('client', '')
    author         = plan.get('author', '')
    doc_date       = plan.get('date', str(date.today()))
    confidentiality = plan.get('confidentiality', 'Internal')

    # Flat sections dict — for keyword-based fallback only
    sections = parse_md_sections(source_text)

    def s(mapping_key: str, *fallback_keys: str) -> str | None:
        """
        Resolve a section-mapping key to content.
        1. If the plan has a heading for this key, return the FULL hierarchical
           content under that heading (all nested sub-headings included).
        2. Otherwise, fall back to a keyword search over heading titles and return
           the full hierarchical content of the first match.
        """
        heading = plan.get('section_mapping', {}).get(mapping_key)
        if heading:
            content = get_section_full(source_text, heading)
            if content:
                return content
        if fallback_keys:
            return find_content_full(source_text, list(fallback_keys))
        return None

    add_doc_header(doc, title, version, status)
    add_page_number_footer(doc)

    # ── COVER PAGE ──────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project:
        p2 = doc.add_paragraph(project)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.size = Pt(14)

    doc.add_paragraph()

    meta = doc.add_table(rows=6, cols=2)
    meta.style = 'Table Grid'
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate([
        ('Document Title',        title),
        ('Project',               project or '—'),
        ('Client / Organisation', client or '—'),
        ('Author',                author or '—'),
        ('Version',               f'v{version} — {status}'),
        ('Date',                  doc_date),
    ]):
        row = meta.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], 'F0F0F0')
        bold_cell(row.cells[0])

    doc.add_paragraph()
    cp = doc.add_paragraph(f'Classification: {confidentiality}')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_page_break()

    # ── DOCUMENT CONTROL ────────────────────────
    doc.add_heading('Document Control', level=1)
    doc.add_heading('Version History', level=2)

    vh = doc.add_table(rows=2, cols=5)
    vh.style = 'Table Grid'

    for i, h_label in enumerate(['Version', 'Date', 'Author', 'Status', 'Description of Changes']):
        cell = vh.rows[0].cells[i]
        cell.text = h_label
        set_cell_bg(cell, 'D4D4D4')
        bold_cell(cell)

    for i, val in enumerate([f'v{version}', doc_date, author or '—', status, 'Initial version']):
        vh.rows[1].cells[i].text = val

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────
    doc.add_heading('Table of Contents', level=1)
    add_toc(doc)
    p = doc.add_paragraph()
    run = p.add_run('Open this document in Microsoft Word and press Ctrl+A → F9 to update the TOC.')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── SECTION 1: INTRODUCTION ─────────────────
    doc.add_heading('1. Introduction', level=1)

    doc.add_heading('1.1 Purpose', level=2)
    c = s('purpose', 'purpose')
    render_md_body(doc, c) if c else tbd(doc, '1.1 Purpose')

    doc.add_heading('1.2 Scope', level=2)
    c = s('scope', 'scope')
    render_md_body(doc, c) if c else tbd(doc, '1.2 Scope')

    doc.add_heading('1.3 Definitions, Acronyms, and Abbreviations', level=2)
    c = s('definitions', 'definitions', 'terms', 'glossary', 'acronyms', 'abbreviations')
    if c:
        render_md_body(doc, c)
    else:
        doc.add_paragraph('Refer to Appendix A (Glossary) and Appendix B (Acronyms and Abbreviations).')

    doc.add_heading('1.4 References', level=2)
    c = s('references', 'references', 'reference documents')
    render_md_body(doc, c) if c else tbd(doc, '1.4 References')

    doc.add_heading('1.5 Document Overview', level=2)
    c = find_content_full(source_text, ['document overview', 'document structure'])
    if c:
        render_md_body(doc, c)
    else:
        doc.add_paragraph(
            'This document follows the IEEE 830-1998 SRS structure. Section 2 provides an overall '
            'system description. Section 3 details functional requirements. Section 4 covers '
            'non-functional requirements. Section 5 describes external interface requirements. '
            'Appendices provide supporting reference material.'
        )

    # ── SECTION 2: OVERALL DESCRIPTION ──────────
    doc.add_heading('2. Overall Description', level=1)

    doc.add_heading('2.1 Product Perspective', level=2)
    c = s('product_perspective', 'product perspective', 'system overview', 'system context',
          'product context', 'background', 'workflow')
    render_md_body(doc, c) if c else tbd(doc, '2.1 Product Perspective')

    doc.add_heading('2.2 Product Functions', level=2)
    c = s('product_functions', 'product functions', 'system functions', 'key features',
          'capabilities', 'functional capabilities')
    render_md_body(doc, c) if c else tbd(doc, '2.2 Product Functions')

    doc.add_heading('2.3 User Classes and Characteristics', level=2)
    c = s('user_classes', 'user classes', 'users', 'stakeholders', 'actors', 'personas')
    render_md_body(doc, c) if c else tbd(doc, '2.3 User Classes and Characteristics')

    doc.add_heading('2.4 Operating Environment', level=2)
    c = s('operating_environment', 'operating environment', 'deployment', 'infrastructure', 'platform')
    render_md_body(doc, c) if c else tbd(doc, '2.4 Operating Environment')

    doc.add_heading('2.5 Assumptions and Dependencies', level=2)
    c = s('assumptions', 'assumptions', 'dependencies', 'assumptions and dependencies')
    render_md_body(doc, c) if c else tbd(doc, '2.5 Assumptions and Dependencies')

    doc.add_heading('2.6 Constraints', level=2)
    c = s('constraints', 'constraints', 'limitations', 'design constraints')
    render_md_body(doc, c) if c else tbd(doc, '2.6 Constraints')

    # ── SECTION 3: FUNCTIONAL REQUIREMENTS ──────
    doc.add_heading('3. Functional Requirements', level=1)
    c = s('functional_requirements')
    if not c:
        # Keyword fallback: look for a heading whose title contains 'functional requirement'
        # but is NOT the document title (skip headings at level 1 that are doc titles)
        for line in source_text.splitlines():
            m = re.match(r'^(#{2,4})\s+(.+)$', line)   # ## or deeper only — skip H1 title
            if m:
                title_text = m.group(2).strip()
                if 'functional requirement' in title_text.lower():
                    content = get_section_full(source_text, title_text)
                    if content:
                        c = content
                        break
    render_md_body(doc, c) if c else tbd(doc, '3 Functional Requirements')

    # ── SECTION 4: NON-FUNCTIONAL REQUIREMENTS ──
    doc.add_heading('4. Non-Functional Requirements', level=1)
    c = s('non_functional_requirements', 'non-functional requirements', 'nfr', 'non functional',
          'quality attributes', 'performance requirements', 'security requirements',
          'technical requirements')
    if c:
        render_md_body(doc, c)
    else:
        for subsec in ['4.1 Performance', '4.2 Security', '4.3 Usability',
                       '4.4 Reliability and Availability', '4.5 Scalability']:
            doc.add_heading(subsec, level=2)
            tbd(doc, subsec)

    # ── SECTION 5: EXTERNAL INTERFACES ──────────
    doc.add_heading('5. External Interface Requirements', level=1)
    c = s('external_interfaces', 'external interface', 'interfaces', 'api', 'integration')
    if c:
        render_md_body(doc, c)
    else:
        subsections = [
            ('5.1 User Interfaces',        ['user interface', 'ui', 'ux']),
            ('5.2 Hardware Interfaces',    ['hardware interface', 'hardware']),
            ('5.3 Software Interfaces',    ['software interface', 'api', 'integration']),
            ('5.4 Communication Interfaces', ['communication', 'network', 'protocol']),
        ]
        for label, keys in subsections:
            doc.add_heading(label, level=2)
            sub = find_content_full(source_text, keys)
            render_md_body(doc, sub) if sub else tbd(doc, label)

    # ── SECTION 6: DATA MIGRATION (optional) ────
    migration_heading = plan.get('section_mapping', {}).get('data_migration')
    if migration_heading:
        migration_content = get_section_full(source_text, migration_heading)
        if migration_content:
            doc.add_heading('6. Data Migration Requirements', level=1)
            render_md_body(doc, migration_content)

    # ── APPENDIX A: GLOSSARY ─────────────────────
    doc.add_page_break()
    doc.add_heading('Appendix A: Glossary', level=1)

    c = s('glossary', 'glossary', 'definitions', 'terms and definitions')
    if c:
        entries = parse_glossary_entries(c)
        if entries:
            add_glossary_table(doc, entries, col_header='Term')
        else:
            render_md_body(doc, c)
    else:
        add_glossary_table(doc, [], col_header='Term')

    # ── APPENDIX B: ACRONYMS ─────────────────────
    doc.add_heading('Appendix B: Acronyms and Abbreviations', level=1)

    c = s('acronyms', 'acronyms', 'abbreviations', 'acronyms and abbreviations')
    if c:
        entries = parse_glossary_entries(c)
        add_glossary_table(doc, entries, col_header='Acronym')
    else:
        # Auto-populate from the built-in lookup; filter to acronyms that appear in the source
        src_upper = source_text.upper()
        entries = [(a, exp) for a, exp in _COMMON_ACRONYMS if a.replace('/', '') in src_upper]
        add_glossary_table(doc, sorted(entries), col_header='Acronym')

    # ── APPENDIX C: OPEN ISSUES ───────────────────
    # Collect content from all gap/open-issue mapping keys in the plan
    appendix_c_parts: list[str] = []
    for gap_key in ('open_issues', 'functional_gaps', 'nfr_gaps', 'migration_gaps', 'gaps'):
        h = plan.get('section_mapping', {}).get(gap_key)
        if h:
            part = get_section_full(source_text, h)
            if part:
                appendix_c_parts.append(part)

    if not appendix_c_parts:
        # Fallback: keyword search for a section that looks like an open-issues list
        fallback = find_content_full(
            source_text,
            ['open issues', 'known issues', 'tbd items', 'risks', 'gaps', 'open questions']
        )
        if fallback:
            appendix_c_parts.append(fallback)

    if appendix_c_parts:
        doc.add_page_break()
        doc.add_heading('Appendix C: Open Issues', level=1)
        for part in appendix_c_parts:
            render_md_body(doc, part)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Document saved: {output_path}")


# ─────────────────────────────────────────────
# CLI entry point
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Generate SRS DOCX from requirements markdown')
    parser.add_argument('--source', required=True, help='Path to source markdown file')
    parser.add_argument('--output', required=True, help='Path for output DOCX file')
    parser.add_argument('--plan',   required=True, help='Path to document plan JSON')
    args = parser.parse_args()

    source = Path(args.source)
    output = Path(args.output)
    plan_file = Path(args.plan)

    if not source.exists():
        sys.exit(f'Source file not found: {source}')
    if not plan_file.exists():
        sys.exit(f'Plan file not found: {plan_file}')

    plan = json.loads(plan_file.read_text(encoding='utf-8'))
    source_text = source.read_text(encoding='utf-8', errors='replace')

    generate(plan, source_text, output)


if __name__ == '__main__':
    main()
